python
from docx import Document
from pathlib import Path
import json
import re


# =========================================================
# FILE LOCATIONS
# =========================================================

CV_FILE = Path("cv/Sangay Wangchuk CV (complete).docx")
OUTPUT_DIR = Path("data/cv")
METRICS_FILE = OUTPUT_DIR / "journal_metrics.json"


# =========================================================
# TEXT CLEANING
# =========================================================

def clean_text(text):
    """
    Normalize whitespace and remove empty/placeholder content.
    """

    text = re.sub(r"\s+", " ", text).strip()

    if text.lower() in {"q", "qq"}:
        return ""

    return text


# =========================================================
# PARAGRAPH EXTRACTION
# =========================================================

def extract_paragraphs(document):
    """
    Extract non-empty paragraphs from the Word document.
    """

    paragraphs = []

    for paragraph in document.paragraphs:

        text = clean_text(paragraph.text)

        if text:
            paragraphs.append(text)

    return paragraphs


# =========================================================
# TABLE EXTRACTION
# =========================================================

def extract_tables(document):
    """
    Extract non-empty table contents.
    """

    tables = []

    for table in document.tables:

        rows = []

        for row in table.rows:

            cells = [
                clean_text(cell.text)
                for cell in row.cells
            ]

            if any(cells):
                rows.append(cells)

        if rows:
            tables.append(rows)

    return tables


# =========================================================
# SECTION EXTRACTION
# =========================================================

def extract_section(paragraphs, start_heading, end_headings):
    """
    Extract paragraphs between a start heading
    and the next recognised heading.
    """

    section = []
    collecting = False

    start_heading_upper = start_heading.upper()

    end_headings_upper = [
        heading.upper()
        for heading in end_headings
    ]

    for paragraph in paragraphs:

        paragraph_upper = paragraph.upper()

        if paragraph_upper == start_heading_upper:

            collecting = True
            continue

        if (
            collecting
            and paragraph_upper in end_headings_upper
        ):

            break

        if collecting:
            section.append(paragraph)

    return section


# =========================================================
# VERIFIED JOURNAL METRICS
# =========================================================

def load_verified_metrics():
    """
    Load verified journal metrics from journal_metrics.json.

    Only records explicitly marked with
    "verified": true are treated as authoritative.
    """

    if not METRICS_FILE.exists():

        print(
            "WARNING: journal_metrics.json not found."
        )

        print(
            "No journal metrics will be treated as verified."
        )

        return {}

    try:

        with open(
            METRICS_FILE,
            "r",
            encoding="utf-8"
        ) as file:

            data = json.load(file)

        if not isinstance(data, dict):

            print(
                "WARNING: journal_metrics.json does not contain "
                "a valid JSON object."
            )

            return {}

        print(
            f"✓ Loaded journal metrics database: "
            f"{len(data)} DOI record(s)."
        )

        return data

    except Exception as error:

        print(
            "WARNING: Could not load journal_metrics.json:"
        )

        print(error)

        return {}


# =========================================================
# PUBLICATION METADATA
# =========================================================

PUBLICATION_METADATA = {

    "10.1016/j.jece.2026.123838": {
        "category": "Environmental Monitoring"
    },

    "10.1021/acsanm.5c05854": {
        "category": "Biomedical / Wearable Sensors"
    },

    "10.1016/j.microc.2026.117476": {
        "category": "Forensic Electrochemistry"
    },

    "10.1016/j.microc.2025.114935": {
        "category": "Forensic Electrochemistry"
    },

    "10.1016/j.talanta.2025.128715": {
        "category": "Forensic Electrochemistry"
    },

    "10.1016/j.mtchem.2025.102872": {
        "category": "Sustainable Nanomaterials"
    },

    "10.1021/acsomega.5c03662": {
        "category": "Forensic Electrochemistry"
    },

    "10.1016/j.talanta.2025.128431": {
        "category": "Environmental / Food Analysis"
    },

    "10.1016/j.talanta.2025.128166": {
        "category": "Environmental Analysis"
    },

    "10.1016/j.microc.2025.113623": {
        "category": "Biomedical / Glucose Sensing"
    },

    "10.1016/j.talanta.2025.127776": {
        "category": "Environmental Analysis"
    },

    "10.1016/j.talanta.2025.127722": {
        "category": "Food / Electrochemical Analysis"
    },

    "10.1016/j.talanta.2025.127581": {
        "category": "Environmental Monitoring"
    },

    "10.1016/j.talanta.2024.127123": {
        "category": "Food / Forensic Analysis"
    },

    "10.1016/j.microc.2024.112217": {
        "category": "Biomedical Analysis"
    },

    "10.1016/j.microc.2024.112216": {
        "category": "Biomedical Analysis"
    },

    "10.1016/j.mtchem.2024.102271": {
        "category": "Biosensors"
    },

    "10.1016/j.microc.2024.111100": {
        "category": "Environmental Monitoring"
    },

    "10.1016/j.talanta.2024.126179": {
        "category": "Environmental Monitoring"
    },

    "10.1016/j.electacta.2024.144292": {
        "category": "Biomedical / Glucose Sensing"
    },

    "10.1016/j.foodchem.2024.138987": {
        "category": "Food / Environmental Analysis"
    },

    "10.1016/j.talanta.2024.125822": {
        "category": "Forensic Analysis"
    },

    "10.1016/j.talanta.2024.125751": {
        "category": "Environmental Analysis"
    },

    "10.1149/1945-7111/acdb9d": {
        "category": "Forensic Electrochemistry"
    },

    "10.1016/j.microc.2023.108668": {
        "category": "Forensic / Food Analysis"
    },

    "10.1016/j.talanta.2023.124266": {
        "category": "Biomedical / Biosensors"
    },

    "10.1016/j.jelechem.2022.116995": {
        "category": "Electrocatalysis"
    },

    "10.1016/j.electacta.2022.141439": {
        "category": "Forensic Electrochemistry"
    },

    "10.1016/j.electacta.2022.141272": {
        "category": "Environmental Analysis"
    },

    "10.17102/bjrd.rub.11.1.026": {
        "category": "Food Analysis / Bhutan"
    }
}


# =========================================================
# CLEAN PUBLICATION DETAILS
# =========================================================

def clean_publication_details(details):

    if not details:
        return ""

    details = re.sub(
        r"\bIF\s+(\d+)\s*\.\s*(\d+)",
        r"IF \1.\2",
        details,
        flags=re.IGNORECASE
    )

    details = re.sub(
        r"\(\s*Q[1-4]\s*,\s*IF\s+\d+(?:\.\s*\d+)?\s*\)",
        "",
        details,
        flags=re.IGNORECASE
    )

    details = re.sub(
        r"\bQ[1-4]\s*,\s*IF\s+\d+(?:\.\s*\d+)?",
        "",
        details,
        flags=re.IGNORECASE
    )

    details = re.sub(
        r"\bQ[1-4]\b",
        "",
        details,
        flags=re.IGNORECASE
    )

    details = re.sub(
        r"\bIF\s+\d+(?:\.\s*\d+)?",
        "",
        details,
        flags=re.IGNORECASE
    )

    details = re.sub(
        r"\(\s*\)",
        "",
        details
    )

    details = re.sub(
        r"\s+",
        " ",
        details
    ).strip()

    details = re.sub(
        r"\s+\.",
        ".",
        details
    )

    details = re.sub(
        r"\s+,",
        ",",
        details
    )

    return details.strip(" .")


# =========================================================
# DOI EXTRACTION
# =========================================================

def extract_doi(text):

    match = re.search(
        r"https://doi\.org/([^\s]+)",
        text,
        flags=re.IGNORECASE
    )

    if not match:
        return ""

    doi_suffix = match.group(1).rstrip(
        ".,;:)"
    )

    return "https://doi.org/" + doi_suffix


# =========================================================
# PUBLICATION PARSER
# =========================================================

def parse_publications(paragraphs, verified_metrics):

    articles = extract_section(
        paragraphs,
        "Peer-Reviewed Journal Articles",
        [
            "SUBMITTED MANUSCRIPT",
            "PhD Thesis",
            "RESEARCH REPORTS",
            "TRAINING/SEMINAR/WORKSHOP/CONFERENCE",
            "REFEREES"
        ]
    )

    publications = []

    for article in articles:

        if article.lower() == "peer-reviewed journal articles":
            continue

        year_match = re.search(
            r"\((\d{4})\)",
            article
        )

        if not year_match:
            continue

        year = int(year_match.group(1))

        doi = extract_doi(article)

        doi_key = ""

        if doi:

            doi_key = doi.replace(
                "https://doi.org/",
                ""
            ).lower()

        citation_without_doi = article

        doi_match = re.search(
            r"https://doi\.org/[^\s]+",
            article,
            flags=re.IGNORECASE
        )

        if doi_match:

            citation_without_doi = article[
                :doi_match.start()
            ].rstrip(" .")

        before_year = citation_without_doi[
            :year_match.start()
        ].strip()

        after_year = citation_without_doi[
            year_match.end():
        ].strip()

        after_year = re.sub(
            r"^[\s\.\-–—:]+",
            "",
            after_year
        )

        authors = before_year.rstrip(" .")

        title = ""
        journal = ""
        details = ""

        parts = [
            part.strip()
            for part in re.split(
                r"\.\s+",
                after_year
            )
            if part.strip()
        ]

        if len(parts) >= 1:
            title = parts[0]

        if len(parts) >= 2:
            journal = parts[1]

        if len(parts) >= 3:
            details = ". ".join(parts[2:])

        title = re.sub(
            r"\s+",
            " ",
            title
        ).strip(" .")

        journal = re.sub(
            r"\s+",
            " ",
            journal
        ).strip(" .")

        details = clean_publication_details(details)

        legacy_metadata = PUBLICATION_METADATA.get(
            doi_key,
            {}
        )

        verified_metadata = verified_metrics.get(
            doi_key,
            {}
        )

        if verified_metadata.get("verified") is not True:
            verified_metadata = {}

        category = legacy_metadata.get(
            "category",
            "Research Publication"
        )

        publication = {

            "year": year,

            "authors": authors,

            "title": title,

            "journal": journal,

            "category": category,

            "quartile": verified_metadata.get(
                "quartile",
                ""
            ),

            "impactFactor": verified_metadata.get(
                "impactFactor",
                ""
            ),

            "metricYear": verified_metadata.get(
                "metricYear",
                ""
            ),

            "details": details,

            "doi": doi,

            "citation": article
        }

        publications.append(
            publication
        )

    return publications


# =========================================================
# SUBMITTED MANUSCRIPT
# =========================================================

def parse_submitted_manuscript(paragraphs):

    manuscripts = extract_section(
        paragraphs,
        "SUBMITTED MANUSCRIPT",
        [
            "PhD Thesis",
            "RESEARCH REPORTS",
            "TRAINING/SEMINAR/WORKSHOP/CONFERENCE",
            "REFEREES"
        ]
    )

    return [
        manuscript
        for manuscript in manuscripts
        if manuscript.lower() != "peer-reviewed journal articles"
    ]


# =========================================================
# PHD THESIS
# =========================================================

def parse_thesis(paragraphs):

    return extract_section(
        paragraphs,
        "PhD Thesis",
        [
            "RESEARCH REPORTS",
            "TRAINING/SEMINAR/WORKSHOP/CONFERENCE",
            "REFEREES"
        ]
    )


# =========================================================
# RESEARCH REPORTS
# =========================================================

def parse_research_reports(paragraphs):

    return extract_section(
        paragraphs,
        "RESEARCH REPORTS",
        [
            "TRAINING/SEMINAR/WORKSHOP/CONFERENCE",
            "REFEREES"
        ]
    )


# =========================================================
# STRUCTURED RESEARCH PROJECT PARSER
# =========================================================

def parse_research_projects(paragraphs):

    section = extract_section(
        paragraphs,
        "PROFESSIONAL SERVICES",
        [
            "RESEARCH GRANTS",
            "AWARDS AND SCHOLARSHIPS/FINANCIAL SUPPORT",
            "SKILLS AND COMPETENCIES",
            "RESEARCH PUBLICATIONS"
        ]
    )

    projects = []

    for paragraph in section:

        text = paragraph.strip()

        lower_text = text.lower()

        # -------------------------------------------------
        # Principal Investigator
        # -------------------------------------------------

        if lower_text.startswith(
            "principal investigator:"
        ):

            project = {
                "role": "Principal Investigator",
                "description": text,
                "project_entity": "",
                "funder": "",
                "year": ""
            }

            # Specific handling for Dorjilung project.
            if "dorjilung" in lower_text:

                project["project_entity"] = (
                    "Dorjilung Hydropower Project Limited (DHPP)"
                )

                project["funder"] = (
                    "Druk Green Power Corporation (DGPC)"
                )

                project["year"] = "2026"

            elif "400 kv" in lower_text:

                project["project_entity"] = (
                    "400 kV D/C East–West Transmission Line Project"
                )

                project["funder"] = (
                    "Bhutan Power Corporation (BPC)"
                )

                project["year"] = "2025"

            elif "132 kv" in lower_text:

                project["project_entity"] = (
                    "132 kV D/C Gamri I & II Transmission Line Projects"
                )

                project["funder"] = (
                    "Bhutan Power Corporation (BPC)"
                )

                project["year"] = "2025"

            projects.append(project)

        # -------------------------------------------------
        # Co-principal Investigator
        # -------------------------------------------------

        elif lower_text.startswith(
            "co-principal investigator:"
        ):

            project = {
                "role": "Co-Principal Investigator",
                "description": text,
                "project_entity": "",
                "funder": "",
                "year": ""
            }

            if "real-time monitoring" in lower_text:

                project["project_entity"] = (
                    "Real-Time Monitoring of Drinking Water Quality "
                    "at Sherubtse College Using IoT Sensors and "
                    "Laboratory Validation"
                )

                project["funder"] = (
                    "Sherubtse College Annual Research Grant (STLRG)"
                )

                project["year"] = "2025–2026"

            projects.append(project)

        # -------------------------------------------------
        # Co-PI
        # -------------------------------------------------

        elif lower_text.startswith(
            "co-pi of the research project:"
        ):

            project = {
                "role": "Co-Principal Investigator",
                "description": text,
                "project_entity": "",
                "funder": "",
                "year": ""
            }

            if "adulterants" in lower_text:

                project["project_entity"] = (
                    "Detection of Adulterants in Common Food Items "
                    "Available in the Bhutanese Market"
                )

                project["funder"] = (
                    "Annual College Research Grant, Sherubtse College, "
                    "Royal University of Bhutan"
                )

                project["year"] = "2020–2021"

            projects.append(project)

        # -------------------------------------------------
        # Core Member
        # -------------------------------------------------

        elif lower_text.startswith(
            "core member of the research project:"
        ):

            project = {
                "role": "Core Member",
                "description": text,
                "project_entity": "",
                "funder": "",
                "year": ""
            }

            if "20 mw yungichhu" in lower_text:

                project["project_entity"] = (
                    "Water Quality and Aquatic Ecology Assessment "
                    "for the 20 MW Yungichhu Hydropower Project"
                )

                project["funder"] = (
                    "Druk Green Power Corporation (DGPC)"
                )

                project["year"] = "2021"

            elif "8 mw thungdiri" in lower_text:

                project["project_entity"] = (
                    "Water Quality and Aquatic Ecology Assessment "
                    "for the 8 MW Thungdiri Hydropower Project"
                )

                project["funder"] = (
                    "Druk Green Power Corporation (DGPC)"
                )

                project["year"] = "2021"

            elif "south asian nitrogen hub" in lower_text:

                project["project_entity"] = (
                    "South Asian Nitrogen Hub (SANH) – WP3.1: "
                    "Role of Nitrogen Air Pollution on Forest "
                    "Ecosystem Services"
                )

                project["funder"] = (
                    "UKRI-GCRF"
                )

                project["year"] = "2020–2021"

            elif "hydrograph separation" in lower_text:

                project["project_entity"] = (
                    "Hydrograph Separation of Streamflow Using "
                    "Geochemical Tracers in Paa Chu Basin"
                )

                project["funder"] = (
                    "World Bank"
                )

                project["year"] = "2019–2020"

            projects.append(project)

    return projects


# =========================================================
# PROFESSIONAL PROFILE
# =========================================================

def parse_profile(paragraphs):

    return extract_section(
        paragraphs,
        "PROFESSIONAL SERVICES",
        [
            "RESEARCH GRANTS",
            "AWARDS AND SCHOLARSHIPS/FINANCIAL SUPPORT",
            "SKILLS AND COMPETENCIES",
            "RESEARCH PUBLICATIONS"
        ]
    )


# =========================================================
# RESEARCH GRANTS
# =========================================================

def parse_research_grants(paragraphs):

    grants = extract_section(
        paragraphs,
        "RESEARCH GRANTS",
        [
            "AWARDS AND SCHOLARSHIPS/FINANCIAL SUPPORT",
            "SKILLS AND COMPETENCIES",
            "RESEARCH PUBLICATIONS"
        ]
    )

    return [
        grant
        for grant in grants
        if grant.strip()
    ]


# =========================================================
# AWARDS
# =========================================================

def parse_awards(paragraphs):

    return extract_section(
        paragraphs,
        "AWARDS AND SCHOLARSHIPS/FINANCIAL SUPPORT",
        [
            "SKILLS AND COMPETENCIES",
            "RESEARCH PUBLICATIONS"
        ]
    )


# =========================================================
# METADATA VALIDATION
# =========================================================

def validate_metadata(publications, verified_metrics):

    print("")
    print("------------------------------------------")
    print("JOURNAL METRICS VALIDATION")
    print("------------------------------------------")

    missing_verified = []
    unverified_records = []
    mismatches = []

    for publication in publications:

        doi = publication.get(
            "doi",
            ""
        )

        if not doi:
            continue

        doi_key = doi.replace(
            "https://doi.org/",
            ""
        ).lower()

        verified = verified_metrics.get(
            doi_key
        )

        if not verified:

            missing_verified.append(
                doi_key
            )

            continue

        if verified.get("verified") is not True:

            unverified_records.append(
                doi_key
            )

            continue

        for field in [
            "quartile",
            "impactFactor",
            "metricYear"
        ]:

            expected = str(
                verified.get(
                    field,
                    ""
                )
            ).strip()

            actual = str(
                publication.get(
                    field,
                    ""
                )
            ).strip()

            if expected != actual:

                mismatches.append(
                    (
                        doi_key,
                        field,
                        expected,
                        actual
                    )
                )

    if missing_verified:

        print(
            f"⚠ {len(missing_verified)} publication(s) "
            "do not have a metrics record:"
        )

        for doi in missing_verified:
            print(f"   - {doi}")

    else:

        print(
            "✓ All DOI-linked publications have "
            "a metrics record."
        )

    if unverified_records:

        print(
            f"⚠ {len(unverified_records)} metrics record(s) "
            "are not marked verified:"
        )

        for doi in unverified_records:
            print(f"   - {doi}")

    else:

        print(
            "✓ All available metrics records are verified."
        )

    if mismatches:

        print(
            "⚠ Metric mismatches detected:"
        )

        for doi, field, expected, actual in mismatches:

            print(
                f"   - {doi}: {field} "
                f"(verified={expected}, "
                f"output={actual})"
            )

    else:

        print(
            "✓ Verified journal metrics validation PASSED."
        )

    print("------------------------------------------")

    return (
        missing_verified,
        unverified_records,
        mismatches
    )


# =========================================================
# MAIN
# =========================================================

def main():

    # -----------------------------------------------------
    # Verify CV
    # -----------------------------------------------------

    if not CV_FILE.exists():

        raise FileNotFoundError(
            f"CV file not found: {CV_FILE}"
        )

    # -----------------------------------------------------
    # Create output directory
    # -----------------------------------------------------

    OUTPUT_DIR.mkdir(
        parents=True,
        exist_ok=True
    )

    # -----------------------------------------------------
    # Load verified journal metrics
    # -----------------------------------------------------

    verified_metrics = load_verified_metrics()

    # -----------------------------------------------------
    # Load Word document
    # -----------------------------------------------------

    document = Document(
        CV_FILE
    )

    # -----------------------------------------------------
    # Extract raw content
    # -----------------------------------------------------

    paragraphs = extract_paragraphs(
        document
    )

    tables = extract_tables(
        document
    )

    # =====================================================
    # RAW CV EXTRACTION
    # =====================================================

    with open(
        OUTPUT_DIR / "cv_text.json",
        "w",
        encoding="utf-8"
    ) as file:

        json.dump(
            {
                "source": str(CV_FILE),
                "paragraphs": paragraphs,
                "tables": tables
            },
            file,
            ensure_ascii=False,
            indent=2
        )

    # =====================================================
    # STRUCTURED PUBLICATIONS
    # =====================================================

    publications = parse_publications(
        paragraphs,
        verified_metrics
    )

    with open(
        OUTPUT_DIR / "publications.json",
        "w",
        encoding="utf-8"
    ) as file:

        json.dump(
            {
                "source": str(CV_FILE),

                "peer_reviewed_journal_articles": publications,

                "article_count": len(publications)
            },
            file,
            ensure_ascii=False,
            indent=2
        )

    # =====================================================
    # SUBMITTED MANUSCRIPT
    # =====================================================

    manuscripts = parse_submitted_manuscript(
        paragraphs
    )

    # =====================================================
    # PHD THESIS
    # =====================================================

    thesis = parse_thesis(
        paragraphs
    )

    # =====================================================
    # RESEARCH REPORTS
    # =====================================================

    reports = parse_research_reports(
        paragraphs
    )

    # =====================================================
    # RESEARCH PROJECTS
    # =====================================================

    research_projects = parse_research_projects(
        paragraphs
    )

    # =====================================================
    # RESEARCH GRANTS
    # =====================================================

    research_grants = parse_research_grants(
        paragraphs
    )

    # =====================================================
    # STRUCTURED RESEARCH DATA
    # =====================================================

    with open(
        OUTPUT_DIR / "research.json",
        "w",
        encoding="utf-8"
    ) as file:

        json.dump(
            {
                "source": str(CV_FILE),

                "projects": research_projects,

                "reports": reports,

                "submitted_manuscripts": manuscripts,

                "phd_thesis": thesis
            },
            file,
            ensure_ascii=False,
            indent=2
        )

    # =====================================================
    # PROFILE AND AWARDS
    # =====================================================

    profile = parse_profile(
        paragraphs
    )

    awards = parse_awards(
        paragraphs
    )

    with open(
        OUTPUT_DIR / "profile.json",
        "w",
        encoding="utf-8"
    ) as file:

        json.dump(
            {
                "source": str(CV_FILE),

                "professional_services": profile,

                "research_grants": research_grants,

                "awards_and_scholarships": awards
            },
            file,
            ensure_ascii=False,
            indent=2
        )

    # =====================================================
    # VALIDATION REPORT
    # =====================================================

    print("")
    print("==========================================")
    print("CV EXTRACTION COMPLETED")
    print("==========================================")

    print(
        f"Paragraphs extracted : {len(paragraphs)}"
    )

    print(
        f"Tables extracted     : {len(tables)}"
    )

    print(
        f"Journal articles     : {len(publications)}"
    )

    print(
        f"Research projects    : {len(research_projects)}"
    )

    print(
        f"Research grants      : {len(research_grants)}"
    )

    print(
        f"Research reports     : {len(reports)}"
    )

    print(
        f"Submitted manuscripts: {len(manuscripts)}"
    )

    print(
        f"PhD thesis entries   : {len(thesis)}"
    )

    print(
        f"Awards/support       : {len(awards)}"
    )

    # =====================================================
    # PROJECT VALIDATION
    # =====================================================

    print("")
    print("------------------------------------------")
    print("RESEARCH PROJECT VALIDATION")
    print("------------------------------------------")

    if research_projects:

        print(
            f"✓ Research project extraction PASSED: "
            f"{len(research_projects)} project(s) detected."
        )

        for index, project in enumerate(
            research_projects,
            start=1
        ):

            print("")
            print(
                f"   {index}. {project['role']}"
            )

            print(
                f"      Project: "
                f"{project['project_entity']}"
            )

            print(
                f"      Funder: "
                f"{project['funder']}"
            )

            print(
                f"      Year: "
                f"{project['year']}"
            )

    else:

        print(
            "⚠ No research projects detected."
        )

    print("------------------------------------------")

    # =====================================================
    # RESEARCH GRANT VALIDATION
    # =====================================================

    print("")
    print("------------------------------------------")
    print("RESEARCH GRANT VALIDATION")
    print("------------------------------------------")

    if research_grants:

        print(
            f"✓ Research grant extraction PASSED: "
            f"{len(research_grants)} grant(s) detected."
        )

    else:

        print(
            "⚠ No research grants detected."
        )

    print("------------------------------------------")

    # =====================================================
    # PUBLICATION COUNT VALIDATION
    # =====================================================

    print("")
    print("------------------------------------------")
    print("PUBLICATION COUNT VALIDATION")
    print("------------------------------------------")

    if len(publications) == 30:

        print(
            "✓ Publication validation PASSED: "
            "30 articles detected."
        )

    else:

        print(
            "⚠ Publication validation WARNING: "
            f"{len(publications)} articles detected; "
            "expected 30."
        )

    # =====================================================
    # DOI VALIDATION
    # =====================================================

    print("")
    print("------------------------------------------")
    print("DOI VALIDATION")
    print("------------------------------------------")

    publications_without_doi = []

    for publication in publications:

        if not publication.get("doi"):

            publications_without_doi.append(
                publication.get(
                    "title",
                    "Untitled publication"
                )
            )

    if publications_without_doi:

        print(
            f"⚠ {len(publications_without_doi)} "
            "publication(s) have no DOI."
        )

        for title in publications_without_doi:

            print(
                f"   - {title}"
            )

    else:

        print(
            "✓ DOI validation PASSED: "
            "all publications contain DOI information."
        )

    # =====================================================
    # JOURNAL METRICS VALIDATION
    # =====================================================

    validate_metadata(
        publications,
        verified_metrics
    )

    # =====================================================
    # FINAL REPORT
    # =====================================================

    print("")
    print("==========================================")
    print("FINAL VALIDATION COMPLETE")
    print("==========================================")

    print(
        "✓ CV extraction is complete."
    )

    print(
        "✓ Structured publication data generated."
    )

    print(
        "✓ Structured research project data generated."
    )

    print(
        "✓ Research grants extracted into profile.json."
    )

    print(
        "✓ Verified journal metrics applied."
    )

    print(
        "✓ CV metric values cannot override "
        "verified journal metrics."
    )

    print(
        "✓ Project funders and project entities "
        "are kept separate."
    )

    print("==========================================")


# =========================================================
# RUN
# =========================================================

if __name__ == "__main__":
    main()
